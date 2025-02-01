from typing import List, Dict, Union, Tuple
import PyPDF2
from docx import Document
import pandas as pd
from io import BytesIO
import requests
from bs4 import BeautifulSoup
import json

class DocumentProcessor:
    def __init__(self):
        self.supported_types = {
            '.pdf': 'PDF Document',
            '.docx': 'Word Document',
            '.xlsx': 'Excel Spreadsheet',
            '.txt': 'Text Document',
            '.html': 'HTML Document',
            '.json': 'JSON Document'
        }
        
    def process_file(self, file_content: bytes, filename: str, content_type: str = None) -> Tuple[str, str]:
        """Main method to process any file type and return text with source information"""
        file_extension = self._get_file_extension(filename)
        
        if file_extension not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        try:
            processed_text = ""
            if file_extension == '.pdf':
                processed_text = self._process_pdf(file_content)
            elif file_extension == '.docx':
                processed_text = self._process_docx(file_content)
            elif file_extension == '.xlsx':
                processed_text = self._process_excel(file_content)
            elif file_extension == '.txt':
                processed_text = self._process_text(file_content)
            elif file_extension == '.html':
                processed_text = self._process_html(file_content)
            elif file_extension == '.json':
                processed_text = self._process_json(file_content)

            # Create source information
            doc_type = self.supported_types[file_extension]
            source_info = f"{doc_type}: {filename}"
            
            return processed_text, source_info

        except Exception as e:
            raise Exception(f"Error processing {filename}: {str(e)}")

    def process_url(self, url: str) -> Tuple[str, str]:
        """Process webpage content with source information"""
        try:
            response = requests.get(url)
            response.raise_for_status()
            processed_text = self._process_html(response.content)
            source_info = f"Web Page: {url}"
            return processed_text, source_info
        except Exception as e:
            raise Exception(f"Error processing URL {url}: {str(e)}")

    def _process_pdf(self, content: bytes) -> str:
        """Extract text from PDF files"""
        text = ""
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
                
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")

    def _process_docx(self, content: bytes) -> str:
        """Extract text from DOCX files"""
        text = ""
        try:
            doc = Document(BytesIO(content))
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Also process tables in the document
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"DOCX processing error: {str(e)}")

    def _process_excel(self, content: bytes) -> str:
        """Extract text from Excel files"""
        try:
            df = pd.read_excel(BytesIO(content))
            text = df.to_string(index=False)
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"Excel processing error: {str(e)}")

    def _process_text(self, content: bytes) -> str:
        """Process plain text files"""
        try:
            return self._clean_text(content.decode('utf-8'))
        except Exception as e:
            raise Exception(f"Text processing error: {str(e)}")

    def _process_html(self, content: bytes) -> str:
        """Extract text from HTML content"""
        try:
            soup = BeautifulSoup(content, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text()
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"HTML processing error: {str(e)}")

    def _process_json(self, content: bytes) -> str:
        """Process JSON files"""
        try:
            data = json.loads(content)
            return json.dumps(data, indent=2)
        except Exception as e:
            raise Exception(f"JSON processing error: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
            
        cleaned = text.strip()
        cleaned = ' '.join(cleaned.split())
        cleaned = '\n'.join(line.strip() for line in cleaned.split('\n') if line.strip())
        
        return cleaned

    def _get_file_extension(self, filename: str) -> str:
        """Get the lowercase file extension from filename"""
        if '.' not in filename:
            return ''
        return filename[filename.rfind('.'):].lower()

# Create a singleton instance
processor = DocumentProcessor()

# Export main functions for easier access
def process_document(file_content: bytes, filename: str, content_type: str = None) -> Tuple[str, str]:
    """Process document and return both text and source information"""
    return processor.process_file(file_content, filename, content_type)

def process_webpage(url: str) -> Tuple[str, str]:
    """Process webpage and return both text and source information"""
    return processor.process_url(url)